"""
Document text extractor for AI question generation.

Supports PDF, DOCX, TXT files and public URLs.
Returns plain text truncated to MAX_CHARS for safe use as an LLM prompt.
"""
from __future__ import annotations

import io
import re
import logging
from typing import Optional

import httpx
from fastapi import UploadFile, HTTPException

logger = logging.getLogger(__name__)

MAX_CHARS = 12_000
MIN_CHARS = 100

_SUPPORTED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
}

_URL_FETCH_TIMEOUT = 15.0
_MAX_URL_BYTES = 5 * 1024 * 1024  # 5 MB for URL content


# ── Helpers ────────────────────────────────────────────────────────────────────

def _truncate(text: str) -> str:
    text = text.strip()
    if len(text) <= MAX_CHARS:
        return text
    return text[:MAX_CHARS] + "\n[... truncated at 12,000 chars]"


def _strip_html(html: str) -> str:
    """Remove HTML tags and collapse whitespace."""
    text = re.sub(r"<script[^>]*>.*?</script>", " ", html, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<style[^>]*>.*?</style>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&[a-z]+;", " ", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _clean_extracted(text: str) -> str:
    """Remove repeated short lines (page numbers, headers/footers)."""
    lines = text.splitlines()
    from collections import Counter
    short_lines = [l.strip() for l in lines if 0 < len(l.strip()) <= 6]
    repeats = {l for l, c in Counter(short_lines).items() if c >= 3}
    cleaned = [l for l in lines if l.strip() not in repeats]
    text = "\n".join(cleaned)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── Extractors ─────────────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF extraction unavailable")

    reader = pypdf.PdfReader(io.BytesIO(data))

    if reader.is_encrypted:
        raise HTTPException(status_code=400, detail="PDF is password-protected. Please provide an unlocked PDF.")

    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue

    text = "\n".join(parts)
    if len(text.strip()) < MIN_CHARS:
        raise HTTPException(
            status_code=422,
            detail="No extractable text found. This appears to be a scanned PDF. Try a text-based PDF."
        )
    return _clean_extracted(text)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX extraction unavailable")

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    text = "\n".join(parts)
    if len(text.strip()) < MIN_CHARS:
        raise HTTPException(status_code=422, detail="Not enough text extracted from this document.")
    return text


def _extract_txt(data: bytes) -> str:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    if len(text.strip()) < MIN_CHARS:
        raise HTTPException(status_code=422, detail="File appears to be empty or too short.")
    return text


# ── Public API ─────────────────────────────────────────────────────────────────

async def extract_from_file(file: UploadFile) -> tuple[str, str]:
    """
    Extract plain text from an uploaded file.
    Returns (text, source_label).
    Raises HTTPException on unsupported type, password protection, or empty content.
    """
    content_type = (file.content_type or "").split(";")[0].strip().lower()
    filename = file.filename or "file"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    # Infer type from extension if content-type is generic
    if content_type in ("application/octet-stream", "binary/octet-stream", ""):
        if ext == "pdf":
            content_type = "application/pdf"
        elif ext == "docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == "txt":
            content_type = "text/plain"

    if content_type not in _SUPPORTED_MIME_TYPES and ext not in ("pdf", "docx", "txt"):
        raise HTTPException(
            status_code=415,
            detail=f"Unsupported file type '{content_type}'. Please upload a PDF, DOCX, or TXT file."
        )

    data = await file.read()

    if content_type == "application/pdf" or ext == "pdf":
        text = _extract_pdf(data)
    elif content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or ext == "docx":
        text = _extract_docx(data)
    else:
        text = _extract_txt(data)

    return _truncate(text), filename


async def extract_from_url(url: str) -> tuple[str, str]:
    """
    Fetch a public URL and extract plain text.
    Returns (text, source_label).
    Raises HTTPException on network errors, access denied, or empty content.
    """
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL must start with http:// or https://")

    try:
        async with httpx.AsyncClient(
            follow_redirects=True,
            timeout=_URL_FETCH_TIMEOUT,
            headers={"User-Agent": "Swaya.me/1.0 (content-extractor)"},
        ) as client:
            resp = await client.get(url)
    except httpx.TimeoutException:
        raise HTTPException(status_code=408, detail="URL request timed out. Check the URL and try again.")
    except httpx.RequestError as e:
        raise HTTPException(status_code=400, detail=f"Could not reach URL: {e}")

    if resp.status_code == 403:
        raise HTTPException(status_code=400, detail="Could not access URL (403 Forbidden). The page may require a login.")
    if resp.status_code == 404:
        raise HTTPException(status_code=400, detail="URL returned 404 Not Found.")
    if resp.status_code >= 400:
        raise HTTPException(status_code=400, detail=f"Could not access URL (HTTP {resp.status_code}).")

    content_type = resp.headers.get("content-type", "").split(";")[0].strip().lower()

    # If the URL points directly to a PDF, extract it as such
    if content_type == "application/pdf":
        if len(resp.content) > _MAX_URL_BYTES:
            raise HTTPException(status_code=413, detail="PDF from URL is too large (max 5 MB).")
        text = _extract_pdf(resp.content)
        label = url.split("/")[-1] or url
        return _truncate(text), label

    if "text" not in content_type and "html" not in content_type:
        raise HTTPException(
            status_code=422,
            detail=f"URL returned unsupported content type '{content_type}'. Only HTML pages and PDFs are supported."
        )

    text = _strip_html(resp.text)
    if len(text.strip()) < MIN_CHARS:
        raise HTTPException(status_code=422, detail="Not enough text found at this URL.")

    label = url.split("//", 1)[-1].split("/")[0]  # domain as label
    return _truncate(text), label

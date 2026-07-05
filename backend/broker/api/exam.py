"""
Exam API — public and authenticated endpoints for exam participation and results.
"""
import io
import re as _re

from fastapi import APIRouter, BackgroundTasks, Body, Depends, HTTPException, Query, Request
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from typing import Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import or_, select

from persistence.database_async import get_async_db
from shared.utils.redis_client import get_redis, RedisClient
from shared.utils.rate_limiter import limiter
from features.quiz.schemas import (
    ExamInfoResponse,
    ExamOtpRequest,
    ExamStartRequest,
    ExamStartResponse,
    ExamAnswerRequest,
    ExamSubmitRequest,
    ExamSubmitResponse,
    ExamResultsResponse,
    ExamPublishResponse,
    AnalyzeResultsRequest,
    ParticipantDetailResponse,
)
from features.quiz import exam_service_async as svc
from shared.exceptions.quiz import QuizNotFoundError, QuizValidationError, InvalidQuizStatusError, ProctoringViolationError
from core.auth.dependencies import get_current_user, CurrentUser, require_admin
from core.ai.gemini_service import analyze_exam_results, generate_interview_sheet, GeminiError

router = APIRouter(tags=["exam"])


@router.post("/e/{slug}/request-otp")
@limiter.limit("10/minute")
async def request_exam_otp(
    request: Request,
    slug: str,
    body: ExamOtpRequest,
    db: AsyncSession = Depends(get_async_db),
    redis: RedisClient = Depends(get_redis),
):
    """Public — send a 6-digit OTP to the participant's email before exam start."""
    try:
        return await svc.request_exam_otp(db, slug, body.display_name, body.email, redis)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.get("/e/{slug}", response_model=ExamInfoResponse)
async def get_exam_info(slug: str, db: AsyncSession = Depends(get_async_db)):
    """Public — get info about an exam (status, dates, question count)."""
    try:
        return await svc.get_exam_info(db, slug)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/e/{slug}/start", response_model=ExamStartResponse)
@limiter.limit("10/minute")
async def start_exam(
    request: Request,
    slug: str,
    body: ExamStartRequest,
    db: AsyncSession = Depends(get_async_db),
    redis: RedisClient = Depends(get_redis),
):
    """Public — start exam; verifies OTP, returns all questions and session_token."""
    try:
        return await svc.start_exam(db, slug, body.display_name, body.email, body.otp, redis)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.post("/e/{slug}/answer")
async def save_answer(slug: str, body: ExamAnswerRequest, db: AsyncSession = Depends(get_async_db)):
    """Public — upsert a single answer; updates last_activity_at."""
    try:
        return await svc.save_answer(
            db, slug, body.session_token, body.question_id, body.selected_option_index,
            text_answer=body.text_answer,
        )
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))


class RunCodeRequest(BaseModel):
    session_token: str
    question_id: int
    language: str
    code: str


import logging as _logging
_run_code_log = _logging.getLogger("exam.run_code")


@router.post("/e/{slug}/run-code")
@limiter.limit("20/minute")
async def run_code(request: Request, slug: str, body: RunCodeRequest, db: AsyncSession = Depends(get_async_db)):
    """Public — participant tests their code; AI evaluates and returns verdict without saving."""
    from core.ai import router as ai_router
    from features.quiz.exam_service_async import _get_active_participant
    from persistence.models.quiz import Quiz, Question, QuestionType
    from sqlalchemy import select as _select

    result = await db.execute(_select(Quiz).filter(Quiz.exam_slug == slug))
    quiz = result.scalar_one_or_none()
    if not quiz:
        raise HTTPException(status_code=404, detail="Exam not found")

    await _get_active_participant(db, quiz, body.session_token)

    q_result = await db.execute(_select(Question).filter(Question.id == body.question_id, Question.quiz_id == quiz.id))
    question = q_result.scalar_one_or_none()
    if not question or question.question_type != QuestionType.CODE:
        raise HTTPException(status_code=404, detail="Question not found or not a code question")

    _run_code_log.warning("run_code called: slug=%s qid=%s lang=%s code_len=%d rubric=%r",
        slug, body.question_id, body.language, len(body.code or ""),
        (question.grading_rubric or "")[:80])
    try:
        ev = await ai_router.evaluate_code(
            language=body.language,
            code=body.code,
            problem_statement=question.text or "",
            grading_rubric=question.grading_rubric or "Evaluate correctness based on the problem statement.",
        )
    except Exception as e:
        _run_code_log.error("run_code ai_router raised: %s %s", type(e).__name__, e)
        raise HTTPException(status_code=503, detail=f"AI evaluation unavailable: {e}")

    _run_code_log.warning("run_code result: verdict=%s explanation=%.80s", ev.get("verdict"), ev.get("explanation"))
    return ev


@router.post("/e/{slug}/submit", response_model=ExamSubmitResponse)
async def submit_exam(
    slug: str,
    body: ExamSubmitRequest,
    db: AsyncSession = Depends(get_async_db),
    redis: RedisClient = Depends(get_redis)
):
    """Public — submit exam; score and return full result."""
    try:
        return await svc.submit_exam(db, slug, body.session_token, redis=redis)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except ProctoringViolationError as e:
        raise HTTPException(status_code=403, detail=str(e))


@router.post("/e/{slug}/result", response_model=ExamSubmitResponse)
async def get_my_result(slug: str, body: ExamSubmitRequest, db: AsyncSession = Depends(get_async_db)):
    """Public — retrieve participant's own score/breakdown after submission."""
    try:
        return await svc.get_my_result(db, slug, body.session_token)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))


@router.get("/quiz/{quiz_id}/exam-results", response_model=ExamResultsResponse)
async def get_exam_results(
    quiz_id: int,
    db: AsyncSession = Depends(get_async_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — full results: leaderboard + per-question analytics."""
    try:
        return await svc.get_exam_results(db, quiz_id, current_user)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except QuizValidationError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.get("/quiz/{quiz_id}/exam-results/participant/{participant_id}", response_model=ParticipantDetailResponse)
async def get_participant_detail(
    quiz_id: int,
    participant_id: int,
    db: AsyncSession = Depends(get_async_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — per-question breakdown for a single participant."""
    try:
        return await svc.get_participant_detail(db, quiz_id, participant_id, current_user)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except QuizValidationError as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/quiz/{quiz_id}/analyze-results")
async def analyze_exam_results_endpoint(
    quiz_id: int,
    body: AnalyzeResultsRequest = Body(default=AnalyzeResultsRequest()),
    db: AsyncSession = Depends(get_async_db),
    current_user: CurrentUser = Depends(require_admin),
):
    """Authenticated host — AI analysis of exam results via Gemini."""
    try:
        results = await svc.get_exam_results(db, quiz_id, current_user)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except QuizValidationError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if results.total_completed == 0:
        raise HTTPException(status_code=400, detail="No completed participants yet — analysis requires at least one submission.")

    try:
        analysis = await analyze_exam_results(results.model_dump(), body.custom_prompt)
    except GeminiError as e:
        raise HTTPException(status_code=503, detail=str(e))

    return {"analysis": analysis}


class SendParticipantEmailsBody(BaseModel):
    sender_name: str | None = None


@router.post("/quiz/{quiz_id}/send-participant-emails")
async def send_participant_emails_now(
    quiz_id: int,
    background_tasks: BackgroundTasks,
    body: SendParticipantEmailsBody = Body(default=SendParticipantEmailsBody()),
    db: AsyncSession = Depends(get_async_db),
    current_user: CurrentUser = Depends(require_admin),
):
    """Authenticated host — queue result emails for all completed participants in the
    background and mark the quiz so the nightly batch skips it. Returns immediately."""
    from persistence.models.quiz import Quiz
    from sqlalchemy import select

    result = await db.execute(select(Quiz).filter(Quiz.id == quiz_id))
    quiz = result.scalar_one_or_none()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")
    if quiz.tenant_id != current_user.tenant_id:
        raise HTTPException(status_code=403, detail="Not authorised")

    # Mark as sent immediately — scheduler won't re-send even if background task is slow
    quiz.exam_participant_emails_sent = True
    await db.commit()

    # Send emails in the background so this request returns instantly
    background_tasks.add_task(
        svc.send_participant_results_emails,
        quiz_id,
        sender_name=body.sender_name or None,
    )

    return {"sent": True, "queued": True}


@router.post("/quizzes/{quiz_id}/publish-exam", response_model=ExamPublishResponse)
async def publish_exam(
    quiz_id: int,
    fresh_start: bool = Query(False),
    db: AsyncSession = Depends(get_async_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — publish exam: generate slug + create session."""
    try:
        return await svc.publish_exam(db, quiz_id, current_user, fresh_start=fresh_start)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except (QuizValidationError, InvalidQuizStatusError) as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/quizzes/{quiz_id}/unpublish-exam")
async def unpublish_exam(
    quiz_id: int,
    db: AsyncSession = Depends(get_async_db),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — unpublish exam (reverts to DRAFT)."""
    try:
        return await svc.unpublish_exam(db, quiz_id, current_user)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except (QuizValidationError, InvalidQuizStatusError) as e:
        raise HTTPException(status_code=400, detail=str(e))


# ── Interview Sheet endpoints ─────────────────────────────────────────────────

_MAX_INTERVIEW_GENERATIONS = 5


class InterviewSheetDownloadRequest(BaseModel):
    sheet: str
    format: str  # "pdf" | "docx" | "md"
    participant_name: str


class InterviewSheetEmailRequest(BaseModel):
    sheet: str
    participant_name: str
    recipient_email: str
    quiz_title: Optional[str] = None


@router.post("/quiz/{quiz_id}/participants/{participant_id}/interview-sheet")
async def generate_interview_sheet_endpoint(
    quiz_id: int,
    participant_id: int,
    db: AsyncSession = Depends(get_async_db),
    redis: RedisClient = Depends(get_redis),
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — generate a Gemini-powered interview sheet for one participant."""
    from persistence.models.quiz import Quiz, FolderShare

    shared_folder_ids = select(FolderShare.folder_id).filter(
        FolderShare.shared_with_user_id == current_user.user_id
    )
    quiz_result = await db.execute(
        select(Quiz).filter(
            Quiz.id == quiz_id,
            or_(
                Quiz.tenant_id == current_user.tenant_id,
                Quiz.folder_id.in_(shared_folder_ids),
            ),
        )
    )
    quiz = quiz_result.scalar_one_or_none()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")

    redis_key = f"interview_sheet_count:{quiz_id}:{participant_id}"
    count = int(await redis.get(redis_key) or 0)
    if count >= _MAX_INTERVIEW_GENERATIONS:
        raise HTTPException(
            status_code=429,
            detail=f"Regeneration limit ({_MAX_INTERVIEW_GENERATIONS}) reached for this participant.",
        )

    try:
        detail = await svc.get_participant_detail(db, quiz_id, participant_id, current_user)
    except QuizNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except QuizValidationError as e:
        raise HTTPException(status_code=400, detail=str(e))

    try:
        sheet = await generate_interview_sheet(detail.model_dump(), quiz.title)
    except GeminiError as e:
        raise HTTPException(status_code=503, detail=str(e))

    await redis.increment(redis_key)

    return {
        "sheet": sheet,
        "participant_name": detail.display_name,
        "generation_count": count + 1,
        "max_generations": _MAX_INTERVIEW_GENERATIONS,
    }


def _name_slug(name: str) -> str:
    return _re.sub(r'[^a-z0-9]+', '_', name.lower()).strip('_') or "participant"


@router.post("/quiz/{quiz_id}/participants/{participant_id}/interview-sheet/download")
async def download_interview_sheet(
    quiz_id: int,
    participant_id: int,
    body: InterviewSheetDownloadRequest,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — download the interview sheet as PDF, Word, or Markdown."""
    slug = _name_slug(body.participant_name)

    if body.format == "md":
        return Response(
            content=body.sheet.encode("utf-8"),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{slug}_interview.md"'},
        )

    if body.format == "docx":
        from docx import Document as _Document

        doc = _Document()
        doc.core_properties.author = "Swaya.me"
        doc.core_properties.subject = f"Interview Sheet — {body.participant_name}"

        for line in body.sheet.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph('')
                continue
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped == '---':
                doc.add_paragraph('─' * 50)
            else:
                para = doc.add_paragraph()
                parts = _re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**') and len(part) > 4:
                        run = para.add_run(part[2:-2])
                        run.bold = True
                    else:
                        para.add_run(part)

        buf = io.BytesIO()
        doc.save(buf)
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{slug}_interview.docx"'},
        )

    if body.format == "pdf":
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=2 * cm, bottomMargin=2 * cm,
        )
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle('IH1', parent=styles['Heading1'], fontSize=18, spaceAfter=12)
        h2 = ParagraphStyle('IH2', parent=styles['Heading2'], fontSize=14, spaceAfter=8, spaceBefore=14)
        h3 = ParagraphStyle('IH3', parent=styles['Heading3'], fontSize=12, spaceAfter=6, spaceBefore=10)
        body_s = ParagraphStyle('IBody', parent=styles['Normal'], fontSize=11, spaceAfter=4, leading=16)

        def _esc(t):
            return t.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        story = []
        for line in body.sheet.split('\n'):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue
            if stripped.startswith('### '):
                story.append(Paragraph(_esc(stripped[4:]), h3))
            elif stripped.startswith('## '):
                story.append(Paragraph(_esc(stripped[3:]), h2))
            elif stripped.startswith('# '):
                story.append(Paragraph(_esc(stripped[2:]), h1))
            elif stripped == '---':
                story.append(Spacer(1, 10))
            else:
                escaped = _re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', _esc(stripped))
                story.append(Paragraph(escaped, body_s))

        doc.build(story)
        return Response(
            content=buf.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{slug}_interview.pdf"'},
        )

    raise HTTPException(status_code=400, detail="format must be 'pdf', 'docx', or 'md'")


@router.post("/quiz/{quiz_id}/participants/{participant_id}/interview-sheet/email")
async def email_interview_sheet(
    quiz_id: int,
    participant_id: int,
    body: InterviewSheetEmailRequest,
    background_tasks: BackgroundTasks,
    current_user: CurrentUser = Depends(get_current_user),
):
    """Authenticated host — email the interview sheet to a recipient."""
    from core.auth.email_service import send_email
    from markdown_it import MarkdownIt

    md = MarkdownIt()
    html_content = md.render(body.sheet)

    quiz_title = body.quiz_title or "Assessment"
    subject = f"Interview Sheet — {body.participant_name} | {quiz_title}"

    sender_name = (
        getattr(current_user, 'display_name', None)
        or getattr(current_user, 'email', None)
        or "Swaya.me"
    )

    html_body = f"""<!DOCTYPE html>
<html>
<head><style>
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; color: #1a1a2e; max-width: 800px; margin: 0 auto; padding: 24px; }}
  h1 {{ color: #1a1a2e; border-bottom: 2px solid #4361ee; padding-bottom: 8px; font-size: 22px; }}
  h2 {{ color: #4361ee; margin-top: 24px; font-size: 17px; }}
  h3 {{ color: #333; font-size: 14px; }}
  p {{ line-height: 1.7; }}
  hr {{ border: none; border-top: 1px solid #e5e7eb; margin: 16px 0; }}
  .footer {{ margin-top: 32px; padding-top: 16px; border-top: 1px solid #e5e7eb; color: #888; font-size: 12px; }}
</style></head>
<body>
{html_content}
<div class="footer">Generated by Swaya.me &mdash; Interview Sheet for {body.participant_name}</div>
</body>
</html>"""

    background_tasks.add_task(
        send_email,
        subject=subject,
        recipients=[body.recipient_email],
        html_body=html_body,
        sender_name=sender_name,
    )

    return {"queued": True}


# ── Certificate endpoints (public — no auth) ───────────────────────────────────

class CertMetaResponse(BaseModel):
    name: str
    quiz_title: str
    score_pct: int
    issued_at: str
    org_name: str
    org_logo_url: Optional[str] = None


async def _get_participant_for_cert(token: str, db: AsyncSession):
    """Look up a completed participant by certificate_token; raise 404 if not found."""
    from persistence.models.quiz import Participant, QuizSession
    from persistence.models.core import Tenant
    from features.quiz.schemas import Quiz

    result = await db.execute(
        select(Participant).where(Participant.certificate_token == token)
    )
    participant = result.scalar_one_or_none()
    if not participant or not participant.completed_at:
        raise HTTPException(status_code=404, detail="Certificate not found.")
    return participant


@router.get("/exam/certificate/{token}")
async def get_certificate_png(token: str, db: AsyncSession = Depends(get_async_db)):
    """
    Public — generate and return the participant's certificate as a PNG.
    Nginx caches this for 1 hour (Cache-Control set here).
    """
    from persistence.models.quiz import Participant, QuizSession
    from persistence.models.core import Tenant
    from persistence.models.quiz import Quiz
    from features.quiz.certificate_service import generate_certificate_png, _resolve_org_logo

    result = await db.execute(
        select(Participant).where(Participant.certificate_token == token)
    )
    participant = result.scalar_one_or_none()
    if not participant or not participant.completed_at:
        raise HTTPException(status_code=404, detail="Certificate not found.")

    # Load the exam session → quiz → tenant
    session_result = await db.execute(
        select(QuizSession).where(QuizSession.id == participant.session_id)
    )
    session = session_result.scalar_one_or_none()
    if not session:
        raise HTTPException(status_code=404, detail="Certificate not found.")

    from sqlalchemy.orm import selectinload as _selectinload
    quiz_result = await db.execute(
        select(Quiz).options(_selectinload(Quiz.questions)).where(Quiz.id == session.quiz_id)
    )
    quiz = quiz_result.scalar_one_or_none()

    tenant_result = await db.execute(select(Tenant).where(Tenant.id == (quiz.tenant_id if quiz else 0)))
    tenant = tenant_result.scalar_one_or_none()

    org_name = (tenant.name if tenant else None) or "Swaya.me"
    org_logo_url = getattr(tenant, "logo_url", None) if tenant else None
    org_logo_path = _resolve_org_logo(org_logo_url)

    # Compute score percentage
    from features.quiz.exam_service_async import _score_participant
    score_pct = await _score_participant(db, participant, quiz)

    png_bytes = generate_certificate_png(
        participant_name=participant.display_name or "Participant",
        quiz_title=quiz.title if quiz else "Assessment",
        score_pct=score_pct,
        issued_at=participant.completed_at,
        org_name=org_name,
        certificate_token=token,
        org_logo_path=org_logo_path,
    )

    return Response(
        content=png_bytes,
        media_type="image/png",
        headers={
            "Content-Disposition": f'attachment; filename="certificate-{token[:8]}.png"',
            "Cache-Control": "public, max-age=3600",
        },
    )


@router.get("/exam/cert-meta/{token}", response_model=CertMetaResponse)
async def get_cert_meta(token: str, db: AsyncSession = Depends(get_async_db)):
    """
    Public — return lightweight certificate metadata for the share landing page.
    """
    from persistence.models.quiz import Participant, QuizSession, Quiz
    from persistence.models.core import Tenant
    from features.quiz.exam_service_async import _score_participant

    result = await db.execute(
        select(Participant).where(Participant.certificate_token == token)
    )
    participant = result.scalar_one_or_none()
    if not participant or not participant.completed_at:
        raise HTTPException(status_code=404, detail="Certificate not found.")

    session_result = await db.execute(
        select(QuizSession).where(QuizSession.id == participant.session_id)
    )
    session = session_result.scalar_one_or_none()

    from sqlalchemy.orm import selectinload as _selectinload2
    quiz_result = await db.execute(
        select(Quiz).options(_selectinload2(Quiz.questions)).where(Quiz.id == (session.quiz_id if session else 0))
    )
    quiz = quiz_result.scalar_one_or_none()

    tenant_result = await db.execute(select(Tenant).where(Tenant.id == (quiz.tenant_id if quiz else 0)))
    tenant = tenant_result.scalar_one_or_none()

    score_pct = await _score_participant(db, participant, quiz)

    return CertMetaResponse(
        name=participant.display_name or "Participant",
        quiz_title=quiz.title if quiz else "Assessment",
        score_pct=score_pct,
        issued_at=participant.completed_at.isoformat(),
        org_name=(tenant.name if tenant else None) or "Swaya.me",
        org_logo_url=getattr(tenant, "logo_url", None) if tenant else None,
    )
